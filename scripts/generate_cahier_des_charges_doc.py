"""
Script pour générer le document Word (.docx) professionnel du Cahier des Charges Final.
"""
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from pathlib import Path
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("ITGate Group — Cahier des charges — Stage d'été 2026")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Plateforme MLOps pour Déploiement IA — ITGate Group | www.itgate-group.com")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(120, 120, 120)

def build_cahier_des_charges():
    doc = Document()
    add_header_footer(doc)
    
    # Couleurs de la charte ITGate
    PRIMARY_COLOR = RGBColor(14, 77, 146)     # Bleu ITGate
    SECONDARY_COLOR = RGBColor(0, 150, 180)   # Teal/Cyan ITGate
    DARK_TEXT = RGBColor(40, 40, 40)
    MUTED_TEXT = RGBColor(100, 100, 100)

    # ─────────────────────────────────────────────────────────────────────────────
    # PAGE DE GARDE
    # ─────────────────────────────────────────────────────────────────────────────
    p_top_space = doc.add_paragraph()
    p_top_space.paragraph_format.space_before = Pt(40)
    
    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_comp = p_company.add_run("ITGATE GROUP")
    r_comp.font.name = 'Calibri'
    r_comp.font.size = Pt(22)
    r_comp.font.bold = True
    r_comp.font.color.rgb = PRIMARY_COLOR
    
    p_subtitle_comp = doc.add_paragraph()
    p_subtitle_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_subtitle_comp.add_run("Solutions Logicielles, Intelligence Artificielle & Cloud\n")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = MUTED_TEXT
    
    p_divider = doc.add_paragraph()
    p_divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_divider.paragraph_format.space_before = Pt(20)
    p_divider.paragraph_format.space_after = Pt(20)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title_main = p_title.add_run("CAHIER DES CHARGES\n")
    r_title_main.font.name = 'Calibri'
    r_title_main.font.size = Pt(26)
    r_title_main.font.bold = True
    r_title_main.font.color.rgb = PRIMARY_COLOR
    
    r_title_sub = p_title.add_run("Plateforme MLOps pour Déploiement & Industrialisation IA")
    r_title_sub.font.name = 'Calibri'
    r_title_sub.font.size = Pt(16)
    r_title_sub.font.bold = True
    r_title_sub.font.color.rgb = SECONDARY_COLOR
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(50)
    
    r_meta = p_meta.add_run("Programme de stages d'été 2026 — Sujet 2\n")
    r_meta.font.name = 'Calibri'
    r_meta.font.size = Pt(13)
    r_meta.font.bold = True
    r_meta.font.color.rgb = DARK_TEXT
    
    r_meta2 = p_meta.add_run("Version Finale 2.0 — Post-Implémentation (Août 2026)\n\n")
    r_meta2.font.name = 'Calibri'
    r_meta2.font.size = Pt(11)
    r_meta2.font.color.rgb = MUTED_TEXT
    
    # Table des métadonnées
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    meta_table.cell(0, 0).paragraphs[0].add_run("Stagiaire :").font.bold = True
    meta_table.cell(0, 1).paragraphs[0].add_run("Hamza Khaled")
    meta_table.cell(1, 0).paragraphs[0].add_run("Encadrant de stage :").font.bold = True
    meta_table.cell(1, 1).paragraphs[0].add_run("Achraf Chehab")
    
    for row in meta_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(80)
    r_foot = p_foot.add_run("www.itgate-group.com")
    r_foot.font.name = 'Calibri'
    r_foot.font.size = Pt(10)
    r_foot.font.color.rgb = MUTED_TEXT
    
    doc.add_page_break()
    
    # ─────────────────────────────────────────────────────────────────────────────
    # SOMMAIRE
    # ─────────────────────────────────────────────────────────────────────────────
    h_som = doc.add_heading("Sommaire", level=1)
    h_som.runs[0].font.color.rgb = PRIMARY_COLOR
    
    sections_list = [
        ("1. Contexte & Objectifs du Projet", "3"),
        ("2. Glossaire des Termes Techniques MLOps & IA", "3"),
        ("3. Cycle de Vie MLOps & Workflow Opérationnel", "4"),
        ("4. Périmètre Fonctionnel & Cas d'Usage Métier", "4"),
        ("5. Architecture Technique en 6 Couches", "6"),
        ("6. Stack Technologique Détaillée & Justifications", "7"),
        ("7. Exigences Non-Fonctionnelles & Critères de Qualité", "8"),
        ("8. Environnement de Développement, Déploiement & CI/CD", "8"),
        ("9. Livrables du Projet & Critères de Réussite", "9"),
    ]
    
    t_som = doc.add_table(rows=len(sections_list), cols=2)
    t_som.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (title, page) in enumerate(sections_list):
        c0 = t_som.cell(idx, 0)
        c1 = t_som.cell(idx, 1)
        r0 = c0.paragraphs[0].add_run(title)
        r0.font.name = 'Calibri'
        r0.font.size = Pt(11)
        r0.font.bold = True if idx in [0, 3, 4, 5, 8] else False
        
        r1 = c1.paragraphs[0].add_run(page)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_margins(c0, top=40, bottom=40, left=50, right=50)
        set_cell_margins(c1, top=40, bottom=40, left=50, right=50)
        
    doc.add_page_break()
    
    # ─────────────────────────────────────────────────────────────────────────────
    # 1. CONTEXTE ET OBJECTIFS
    # ─────────────────────────────────────────────────────────────────────────────
    h1 = doc.add_heading("1. Contexte & Objectifs du Projet", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph()
    p.add_run("ITGate Group ").bold = True
    p.add_run("est une entreprise technologique innovante spécialisée dans le développement logiciel, l'Intelligence Artificielle, le Cloud, le DevOps et les architectures SaaS. Ce document de spécification cadre le projet de stage d'été 2026 intitulé ")
    p.add_run("« Plateforme MLOps pour Déploiement & Industrialisation IA »").bold = True
    p.add_run(" (Sujet 2).\n\n")
    p.add_run("Le passage d'un modèle d'IA du laboratoire (Notebook) à la production d'entreprise nécessite une rigueur d'ingénierie logicielle avancée : reproductibilité, versionnement, monitoring de la dérive des données (Data Drift), résilience d'inférence et automatisation CI/CD & GitOps.\n\n")
    p.add_run("Objectifs majeurs du projet :").bold = True
    
    doc.add_paragraph("Comprendre et implémenter l'intégralité du cycle de vie MLOps d'un modèle d'IA en production.", style='List Bullet')
    doc.add_paragraph("Résoudre un cas d'usage économique réel pour ITGate : la prévision du Chiffre d'Affaires par Séries Temporelles Multi-variées.", style='List Bullet')
    doc.add_paragraph("Intégrer des capacités d'IA Générative et NLP d'entreprise : moteur RAG documentaire (Groq + FAISS), extraction structurée et classification de documents.", style='List Bullet')
    doc.add_paragraph("Implémenter la supervision continue de la qualité de la donnée via la détection automatique de Data Drift.", style='List Bullet')
    doc.add_paragraph("Déployer une architecture scalable conteneurisée (Docker, Kubernetes HPA) avec observabilité (Prometheus, Grafana) et livraison continue GitOps (ArgoCD, GitHub Actions).", style='List Bullet')
    doc.add_paragraph("Fournir une console web de pilotage ergonomique et réactive en React avec la charte officielle ITGate.", style='List Bullet')

    # ─────────────────────────────────────────────────────────────────────────────
    # 2. GLOSSAIRE DES TERMES TECHNIQUES
    # ─────────────────────────────────────────────────────────────────────────────
    h2 = doc.add_heading("2. Glossaire des Termes Techniques MLOps & IA", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR
    
    terms = [
        ("MLOps (Machine Learning Operations)", "Ensemble de pratiques et d'outils unifiant le développement ML (Data Science) et les opérations logicielles (DevOps) pour fiabiliser et accélérer le déploiement continu des modèles."),
        ("Experiment Tracking (MLflow)", "Historisation systématique des hyperparamètres, des métriques d'évaluation (R², RMSE, MAE) et des artefacts générés lors de chaque entraînement."),
        ("Model Registry", "Répertoire centralisé et versionné gérant le cycle de vie formel des modèles validés (stages : Staging, Production, Archived)."),
        ("Model Serving (FastAPI)", "Mise à disposition du modèle entraîné sous la forme d'un service d'API REST haute performance accessible en temps réel."),
        ("Séries Temporelles Multi-variées", "Modélisation prédictive exploitant simultanément l'historique temporel passé (Lag Features) et des indicateurs d'activité métier (ingénieurs, projets, contrats)."),
        ("Data Drift (Dérive des données)", "Altération statistique de la distribution des données reçues en production par rapport à la distribution d'apprentissage, mesurée via le Z-score normalisé."),
        ("RAG (Retrieval-Augmented Generation)", "Architecture combinant la recherche sémantique locale dans une base vectorielle (FAISS) et la puissance de synthèse d'un LLM (Groq) pour répondre précisément sur des documents privés."),
        ("GitOps (ArgoCD)", "Pratique opérationnelle où l'état désiré de l'infrastructure Kubernetes est décrit de façon déclarative dans un dépôt Git et synchronisé automatiquement.")
    ]
    
    for term, definition in terms:
        p = doc.add_paragraph()
        r = p.add_run(f"• {term} : ")
        r.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        p.add_run(definition)

    # ─────────────────────────────────────────────────────────────────────────────
    # 3. CYCLE DE VIE MLOPS
    # ─────────────────────────────────────────────────────────────────────────────
    h3 = doc.add_heading("3. Cycle de Vie MLOps & Workflow Opérationnel", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph("Le cycle MLOps mis en œuvre sur la plateforme est itératif et interconnecté en 7 phases distinctes :")
    
    lifecycle_steps = [
        ("1. Préparation & Ingestion des Données", "Consolidation du dataset multi-varié ITGate et calcul du profil statistique de référence (Baseline JSON)."),
        ("2. Entraînement & Experiment Tracking", "Entraînement du modèle RandomForest avec suivi exhaustif des paramètres et métriques sous MLflow."),
        ("3. Évaluation & Promotion dans le Model Registry", "Validation des performances et promotion automatique du modèle dans le registre MLflow au stage Production."),
        ("4. Empaquetage Multi-stage Docker", "Compilation des interfaces React et empaquetage du runtime Python dans une image conteneurisée sécurisée et légère."),
        ("5. Déploiement & Orchestration Kubernetes", "Instanciation du service avec sondes Liveness/Readiness et auto-scaling horizontal (HPA)."),
        ("6. Monitoring & Détection de Data Drift", "Scraping temps réel des métriques Prometheus et alerte sur les déviations statistiques du trafic entrant."),
        ("7. Intégration Continue & GitOps", "Validation continue des tests via GitHub Actions et réconciliation de déploiement automatique avec ArgoCD.")
    ]
    for step, desc in lifecycle_steps:
        p = doc.add_paragraph()
        r = p.add_run(f"{step} : ")
        r.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────────
    # 4. PÉRIMÈTRE FONCTIONNEL & CAS D'USAGE MÉTIER
    # ─────────────────────────────────────────────────────────────────────────────
    h4 = doc.add_heading("4. Périmètre Fonctionnel & Cas d'Usage Métier", level=1)
    h4.runs[0].font.color.rgb = PRIMARY_COLOR
    
    # 4.1 Time series
    h4_1 = doc.add_heading("4.1 Prévision du Chiffre d'Affaires ITGate (Séries Temporelles Multi-variées)", level=2)
    h4_1.runs[0].font.color.rgb = SECONDARY_COLOR
    doc.add_paragraph("La plateforme intègre un modèle prédictif dédié à la direction financière d'ITGate Group, permettant d'anticiper le Chiffre d'Affaires du mois suivant (M+1).")
    
    ts_table = doc.add_table(rows=7, cols=2)
    ts_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ts_data = [
        ("Dataset d'entraînement", "data/raw/itgate_revenue_multivariate.csv (48 mois d'historique)"),
        ("Algorithme", "RandomForestRegressor (Scikit-Learn)"),
        ("Features Métier", "num_engineers, active_projects, avg_contract_value"),
        ("Features Temporelles (Lags)", "revenue_lag_1 (M-1), revenue_lag_2 (M-2), revenue_lag_3 (M-3)"),
        ("Model Registry", "Enregistré sous 'ITGate_Revenue_Model' (Stage : Production)"),
        ("Endpoint d'inférence", "POST /predict (Chargement dynamique via MLflow Registry)"),
        ("Métriques de performance", "R² Score > 0.85, RMSE, MAE suivis dans MLflow")
    ]
    for idx, (k, v) in enumerate(ts_data):
        c0 = ts_table.cell(idx, 0)
        c1 = ts_table.cell(idx, 1)
        r0 = c0.paragraphs[0].add_run(k)
        r0.font.bold = True
        c1.paragraphs[0].add_run(v)
        set_cell_margins(c0, top=50, bottom=50, left=80, right=80)
        set_cell_margins(c1, top=50, bottom=50, left=80, right=80)
        if idx % 2 == 0:
            set_cell_background(c0, "F0F4F8")
            set_cell_background(c1, "F0F4F8")

    # 4.2 NLP & RAG
    h4_2 = doc.add_heading("4.2 Services NLP & IA Générative (Groq Cloud + FAISS Local)", level=2)
    h4_2.runs[0].font.color.rgb = SECONDARY_COLOR
    doc.add_paragraph("La plateforme expose trois endpoints IA avancés exploitant l'API Groq (modèles LLaMA) couplée à un stockage vectoriel local :")
    
    doc.add_paragraph("Moteur RAG Documentaire (POST /ask) : Permet aux collaborateurs d'interroger la base documentaire interne ITGate. Les textes sont découpés en chunks, vectorisés localement avec sentence-transformers/all-MiniLM-L6-v2 et indexés sous FAISS. Groq génère la réponse finale avec citations des sources. Un fallback transparent assure le retour des sources locales en cas d'indisponibilité du LLM.", style='List Bullet')
    doc.add_paragraph("Extraction Intelligente (POST /extract) : Analyse de documents non-structurés (factures, contrats) avec extraction des entités clés (fournisseur, date, montant) et validation stricte par schéma Pydantic.", style='List Bullet')
    doc.add_paragraph("Classification Automatique (POST /classify) : Catégorisation zero-shot de textes (Facture, CV, Contrat, Rapport) avec filtrage des hallucinations.", style='List Bullet')

    # 4.3 Data drift
    h4_3 = doc.add_heading("4.3 Détection de Data Drift en Temps Réel", level=2)
    h4_3.runs[0].font.color.rgb = SECONDARY_COLOR
    doc.add_paragraph("Pour prévenir l'obsolescence silencieuse du modèle, le module src/drift.py évalue la dérive des données reçues sur POST /predict en calculant le Z-Score par rapport au profil de référence (drift_baseline.json). L'état est restitué via l'endpoint GET /drift et la métrique Prometheus mlops_data_drift_score.")

    # 4.4 Sécurité & Frontend
    h4_4 = doc.add_heading("4.4 Sécurité JWT, Base d'Audit & Console Web React", level=2)
    h4_4.runs[0].font.color.rgb = SECONDARY_COLOR
    doc.add_paragraph("• Sécurité OAuth2/JWT : Authentification par token avec hachage sécurisé pbkdf2_sha256.\n• Traçabilité SQLite/SQLAlchemy : Enregistrement de chaque requête API (latence, prédiction, statut) exposé sur GET /history.\n• Console Web React Glassmorphism : Dashboard interactif intégrant le logo officiel ITGate, visualisation SVG de séries temporelles (historique + projection M+1) et widgets de statut Data Drift.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────────
    # 5. ARCHITECTURE TECHNIQUE
    # ─────────────────────────────────────────────────────────────────────────────
    h5 = doc.add_heading("5. Architecture Technique en 6 Couches", level=1)
    h5.runs[0].font.color.rgb = PRIMARY_COLOR
    
    arch_layers = [
        ("1. Couche ML & NLP (Tracking & Ingestion)", "Entraînement RandomForest (Time Series Lags + Métier ITGate) avec MLflow Tracking (R², RMSE) et Model Registry (Stage Production). Vectorisation Sentence-Transformers + FAISS."),
        ("2. Couche Serving & Backend (FastAPI)", "Exposition des routes /predict, /ask, /extract, /classify, /drift, /history, /metrics et des sondes Kubernetes /health/live et /health/ready. Sécurité JWT et persistance SQLite."),
        ("3. Couche Empaquetage (Docker Multi-stage)", "Construction d'une image conteneur optimisée combinant le build des assets statiques React (Node 22) et l'exécution du serveur FastAPI (Python 3.11-slim) sous utilisateur non-root."),
        ("4. Couche Orchestration (Kubernetes)", "Déploiement managé via Deployment, Service, ConfigMaps, Secrets, Liveness/Readiness probes et Horizontal Pod Autoscaler (HPA basé sur le CPU)."),
        ("5. Couche Observabilité (Prometheus & Grafana)", "Collecte des métriques HTTP, des temps de réponse LLM et du score de Data Drift par Prometheus. Dashboard Grafana pré-provisionné (mlops-platform.json)."),
        ("6. Couche Automatisation & GitOps", "Pipeline GitHub Actions (tests unitaires pytest, validation des artefacts, build Docker, publication GHCR) et synchronisation continue de cluster avec ArgoCD.")
    ]
    
    for title, desc in arch_layers:
        p = doc.add_paragraph()
        r = p.add_run(f"■ {title}\n")
        r.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        p.add_run(desc)

    # ─────────────────────────────────────────────────────────────────────────────
    # 6. STACK TECHNOLOGIQUE
    # ─────────────────────────────────────────────────────────────────────────────
    h6 = doc.add_heading("6. Stack Technologique Détaillée & Justifications", level=1)
    h6.runs[0].font.color.rgb = PRIMARY_COLOR
    
    stack_table = doc.add_table(rows=12, cols=3)
    stack_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Composant", "Technologie", "Rôle & Justification"]
    for i, h in enumerate(headers):
        c = stack_table.cell(0, i)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "0E4D92")
        set_cell_margins(c, top=60, bottom=60, left=80, right=80)
        
    stack_items = [
        ("Machine Learning", "Scikit-Learn", "Modélisation robuste par Random Forest pour données tabulaires et temporelles."),
        ("Experiment Tracking", "MLflow (SQLite Backend)", "Traçabilité des runs, stockage des artefacts et gouvernance via le Model Registry."),
        ("Embeddings & Vector Store", "Sentence-Transformers + FAISS", "Vectorisation sémantique locale haute vitesse sans dépendance cloud."),
        ("Génération IA (LLM)", "Groq Cloud API (LLaMA)", "Inférence ultra-rapide (<300ms) et tier de développement gratuit."),
        ("API Backend", "FastAPI + Uvicorn + Pydantic", "Framework asynchrone performant avec validation stricte et Swagger UI natif."),
        ("Persistance & Sécurité", "SQLAlchemy + SQLite + PyJWT", "Stockage de l'historique d'inférence et contrôle d'accès sécurisé par jetons."),
        ("Interface Frontend", "React 18 + Vite + CSS", "Dashboard moderne Glassmorphism avec graphiques SVG interactifs."),
        ("Conteneurisation", "Docker (Multi-stage)", "Image conteneur unifiée, reproductible et légère (<600 Mo)."),
        ("Orchestration", "Kubernetes (Minikube/Docker)", "Gestion de la résilience, des sondes de vie et de l'autoscaling horizontal."),
        ("Monitoring & Alerting", "Prometheus + Grafana", "Supervision technique et métier (Data Drift, latence LLM, trafic HTTP)."),
        ("CI/CD & GitOps", "GitHub Actions + ArgoCD", "Automatisation des tests, packaging GHCR et déploiement déclaratif continu.")
    ]
    
    for idx, (comp, tech, just) in enumerate(stack_items):
        row_idx = idx + 1
        c0 = stack_table.cell(row_idx, 0)
        c1 = stack_table.cell(row_idx, 1)
        c2 = stack_table.cell(row_idx, 2)
        c0.paragraphs[0].add_run(comp).font.bold = True
        c1.paragraphs[0].add_run(tech)
        c2.paragraphs[0].add_run(just)
        for c in [c0, c1, c2]:
            set_cell_margins(c, top=40, bottom=40, left=60, right=60)
            if row_idx % 2 == 0:
                set_cell_background(c, "F7F9FB")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────────
    # 7. EXIGENCES NON-FONCTIONNELLES
    # ─────────────────────────────────────────────────────────────────────────────
    h7 = doc.add_heading("7. Exigences Non-Fonctionnelles & Critères de Qualité", level=1)
    h7.runs[0].font.color.rgb = PRIMARY_COLOR
    
    doc.add_paragraph("• Performance d'Inférence : Le service /predict doit répondre en moins de 15 ms en charge nominale ; les réponses RAG et IA doivent être restituées en moins de 500 ms.")
    doc.add_paragraph("• Résilience & Haute Disponibilité : Séparation stricte entre la sonde de liveness (/health/live) et la sonde de readiness (/health/ready) garantissant qu'aucune requête n'est routée avant le chargement complet du modèle MLflow.")
    doc.add_paragraph("• Dégradation Gracieuse (Fallback) : En cas de défaillance réseau ou de quota avec Groq, le RAG bascule automatiquement sur les fragments locaux FAISS sans générer d'erreur HTTP 500.")
    doc.add_paragraph("• Sécurité & Confidentialité : Isolation des secrets via variables d'environnement (.env) et Kubernetes Secrets. Aucun identifiant codé en dur.")
    doc.add_paragraph("• Qualité & Maintenabilité : Suite de tests automatisée (unitaires, mocks d'API, validation d'artefacts) garantissant une non-régression à chaque commit.")

    # ─────────────────────────────────────────────────────────────────────────────
    # 8. ENVIRONNEMENT DE DÉVELOPPEMENT & CI/CD
    # ─────────────────────────────────────────────────────────────────────────────
    h8 = doc.add_heading("8. Environnement de Développement, Déploiement & CI/CD", level=1)
    h8.runs[0].font.color.rgb = PRIMARY_COLOR
    
    doc.add_paragraph("Environnement de développement :").bold = True
    doc.add_paragraph("Python 3.11+ avec environnement virtuel dédié (venv).", style='List Bullet')
    doc.add_paragraph("Node.js 22 LTS & npm pour le développement du frontend React.", style='List Bullet')
    doc.add_paragraph("Docker Desktop avec Kubernetes activé pour l'exécution locale.", style='List Bullet')
    doc.add_paragraph("MLflow Tracking Server avec base de données SQLite locale (mlflow.db).", style='List Bullet')
    
    doc.add_paragraph("\nPipeline d'Intégration & Déploiement Continus :").bold = True
    doc.add_paragraph("CI GitHub Actions : Exécution automatique de pytest, vérification des manifestes, build Docker multi-stage et publication vers GitHub Container Registry (ghcr.io).", style='List Bullet')
    doc.add_paragraph("Déploiement GitOps ArgoCD : Manifeste k8s/argocd/application.yaml synchronisant l'état du cluster Kubernetes avec les manifestes du dépôt Git.", style='List Bullet')

    # ─────────────────────────────────────────────────────────────────────────────
    # 9. LIVRABLES DU PROJET & CRITÈRES DE RÉUSSITE
    # ─────────────────────────────────────────────────────────────────────────────
    h9 = doc.add_heading("9. Livrables du Projet & Critères de Réussite", level=1)
    h9.runs[0].font.color.rgb = PRIMARY_COLOR
    
    doc.add_paragraph("Livrables Finaux du Projet :").bold = True
    doc.add_paragraph("1. Code source complet et versionné sur GitHub (hamzakh86/mlops-plateforme).", style='List Bullet')
    doc.add_paragraph("2. Pipeline d'entraînement et Model Registry MLflow (src/train.py, ITGate_Revenue_Model en stage Production).", style='List Bullet')
    doc.add_paragraph("3. API FastAPI d'inférence sécurisée avec supervision Data Drift (src/serve.py, src/drift.py).", style='List Bullet')
    doc.add_paragraph("4. Console Web React intégrée et brandée ITGate Group (frontend/).", style='List Bullet')
    doc.add_paragraph("5. Image Docker multi-stage et manifestes Kubernetes (k8s/deployment.yaml, hpa.yaml, argocd/).", style='List Bullet')
    doc.add_paragraph("6. Stack de monitoring automatisée (Prometheus & dashboard Grafana pré-configuré).", style='List Bullet')
    doc.add_paragraph("7. Notebook interactif de démonstration et soutenance (notebooks/demo_itgate.ipynb).", style='List Bullet')
    doc.add_paragraph("8. Ce document officiel de Cahier des Charges (Version Finale 2.0).", style='List Bullet')
    doc.add_paragraph("9. Rapport de stage professionnel et présentation de soutenance.", style='List Bullet')
    
    doc.add_paragraph("\nCritères de validation validés :").bold = True
    doc.add_paragraph("✔ Modèle entraîné, suivi sous MLflow et déployé dynamiquement en Production.", style='List Bullet')
    doc.add_paragraph("✔ Dérive des données (Data Drift) monitorée en continu via Z-score et Prometheus.", style='List Bullet')
    doc.add_paragraph("✔ Services IA (RAG, Extraction, Classification) opérationnels et résilients.", style='List Bullet')
    doc.add_paragraph("✔ Pipeline CI/CD fonctionnel sans intervention manuelle.", style='List Bullet')
    doc.add_paragraph("✔ Interface utilisateur claire reflétant l'état de la plateforme en temps réel.", style='List Bullet')

    # Sauvegarde du document
    output_docx_path = Path("docs/Cahier_des_Charges_MLOps_ITGate_Final.docx")
    output_doc_path = Path("docs/Cahier_des_Charges_MLOps_ITGate_Final.doc")
    
    doc.save(str(output_docx_path))
    # Copie également en .doc pour compatibilité exacte avec la demande de l'utilisateur
    doc.save(str(output_doc_path))
    
    print(f"[OK] Document généré avec succès : {output_docx_path.resolve()}")
    print(f"[OK] Document généré avec succès : {output_doc_path.resolve()}")

if __name__ == "__main__":
    build_cahier_des_charges()
